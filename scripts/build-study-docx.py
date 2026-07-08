#!/usr/bin/env python3
"""Convert study guide markdown files to Word (.docx) for Google Drive."""

from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

GUIDES_DIR = Path(__file__).resolve().parent.parent / "docs" / "study-guides"
OUTPUT_DIR = GUIDES_DIR / "word"


def add_formatted_run(paragraph, text: str) -> None:
    """Add text with **bold** segments."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    in_code = False
    code_lines: list[str] = []

    lines = md_path.read_text(encoding="utf-8").splitlines()

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not stripped:
            continue

        if stripped == "---":
            doc.add_paragraph()
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("|") and "---" not in stripped:
            p = doc.add_paragraph()
            add_formatted_run(p, stripped)
            p.paragraph_format.left_indent = Inches(0.25)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_run(p, stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_formatted_run(p, re.sub(r"^\d+\.\s", "", stripped))
        else:
            p = doc.add_paragraph()
            add_formatted_run(p, stripped)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Created: {docx_path.name}")


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    md_files = sorted(GUIDES_DIR.glob("*.md"))
    if not md_files:
        print("No markdown files found.")
        return
    for md_path in md_files:
        docx_name = md_path.stem + ".docx"
        md_to_docx(md_path, OUTPUT_DIR / docx_name)
    print(f"\nDone. {len(md_files)} Word files in:\n  {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
